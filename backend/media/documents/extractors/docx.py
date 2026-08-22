"""Bounded DOCX text extractor for the Document Pulper."""

from __future__ import annotations

import io
import logging

from media.documents.security import DocumentSecurityError, validate_office_container

logger = logging.getLogger(__name__)
_MAX_TEXT_CHARS = 1_000_000


def extract_docx(file_bytes: bytes) -> str:
    """Extract bounded text from a validated DOCX container."""
    validate_office_container(file_bytes, "docx")
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        total_chars = 0

        def append_part(text: str) -> None:
            nonlocal total_chars
            if not text:
                return
            total_chars += len(text)
            if total_chars > _MAX_TEXT_CHARS:
                raise DocumentSecurityError("docx_text_limit")
            parts.append(text)

        for para in doc.paragraphs:
            append_part(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                append_part("\t".join(cell.text.strip() for cell in row.cells).strip())

        return "\n".join(parts)

    except DocumentSecurityError:
        raise
    except ImportError:
        logger.error("docx_extractor: python-docx unavailable")
        return ""
    except Exception as exc:
        logger.warning(
            "docx_extractor: failed error_class=%s",
            type(exc).__name__,
        )
        return ""
